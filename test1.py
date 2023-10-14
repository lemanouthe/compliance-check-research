import re
import json
import os
import PyPDF2
from docx import Document
from src import (
    get_confidence_factor,
    get_requirement_type,
    get_requirement_domain,
    get_rfp_domain,
    identify_relationships,
)

from requirementsType.bbc import auto_categorization as types
from requirementsDomain.bbc import auto_categorization as domain
from rfp_domain.bbc import auto_categorization as rfp_domain

import nltk

# categorizer = get_requirement.RequirementCategorizer()


class RequirementsExtractor:
    def __init__(self, file_path):
        self.file_path = file_path
        self.text = ""
        self.paragraphs = ""
        self.cleaned_text = ""
        self.sentences = ""
        self.background_subsection = ""
        self.all_requirements_data = []
        self.extracted_requirements_list = []
        self.segmented_text = []

    def clean_text(self):
        # Supprimer les caractères spéciaux, les sauts de ligne inutiles et les espaces en double
        self.cleaned_text = re.sub(r"[^\w\s]", "", self.text)
        self.cleaned_text = re.sub(r"\n+", "\n", self.cleaned_text)
        self.cleaned_text = re.sub(r"\s+", " ", self.cleaned_text)

        self.cleaned_text.strip()
        
        return self.cleaned_text

    def segment_sentences(self, text):
        # Segmenter le texte en phrases
        sentences = re.findall(r"[^.!?]+[.!?]", text)
        # Tokenisation
        # sentences = nltk.sent_tokenize(text)
        return sentences

    def segment_paragraphs(self):
        # Remplacer les sauts de ligne par des espaces pour nettoyer les paragraphes
        self.cleaned_text = re.sub(r"[^\x00-\x7f]", "", self.text)
        self.cleaned_text = re.sub(r"\n", " ", self.cleaned_text)
        self.cleaned_text = re.sub(r"\d", "", self.cleaned_text)
        self.cleaned_text = re.sub(r"–", "", self.cleaned_text).strip()
        # Segmenter le texte en paragraphes
        self.paragraphs = re.split(r"\n\n+", self.cleaned_text)
        return self.paragraphs

    def is_requirement(self, sentence):
        # Définissez ici vos critères pour identifier une phrase comme étant une exigence
        # Par exemple, vous pouvez vérifier si la phrase contient certains mots clés ou motifs spécifiques.
        self.keywords = [
            "shall",
            "must",
            "must be able to",
            "should",
            "have to",
            "need to",
            "review",
            "edit",
            "update",
            "ensure",
            "prepare and submit",
            "handle",
        ]  # Liste de mots clés à rechercher

        # return any(keyword in sentence.lower() for keyword in keywords) and sentence.strip().endswith('.')
        return any(keyword in sentence.lower() for keyword in self.keywords)

    def read_pdf(self):
        with open(self.file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                self.text += page.extractText()

    def read_docx(self):
        doc = Document(self.file_path)
        for paragraph in doc.paragraphs:
            self.text += paragraph.text + "\n"

    def read_txt(self):
        with open(self.file_path, "r") as txt_file:
            self.text = txt_file.read()

    def extract_background_subsection(self):
        # Define the possible start markers for the "Background" subsection
        start_markers = [
            "1. Background",
            "1.1. Background.",
            "1.0 BACKGROUND",
        ]

        # Define the possible end markers for the subsection
        end_markers = [
            "2.",
            "1.2.",
            "2.0",
        ]

        # Initialize variables to store the start and end positions
        start_index = None
        end_index = None

        # Find the starting position using any of the possible start markers
        for marker in start_markers:
            start_index = self.text.find(marker)
            if start_index != -1:
                break  # Exit the loop if a valid start marker is found

        # Find the ending position using any of the possible end markers
        for marker in end_markers:
            if start_index is not None:
                end_index = self.text.find(marker, start_index)
                if end_index != -1:
                    break  # Exit the loop if a valid end marker is found

        if start_index is not None and end_index != -1:
            # Extract the "Background" subsection
            self.background_subsection = self.text[start_index:end_index].strip()
            # print(self.background_subsection)
            return self.background_subsection
        else:
            return None

    def extract_requirements(self):
        for req in self.segmented_text:
            # Créez un dictionnaire représentant une exigence
            if self.is_requirement(req):
                requirement_data = {
                    "requirement_id": len(self.extracted_requirements_list) + 1,
                    "original_RFP_context": os.path.basename(self.file_path),
                    "extracted_requirements": req.strip(),
                    "requirement_type": types.classify(
                        req
                    ),  # get_requirement_type(req),
                    "RFP_domain": rfp_domain.classify(
                        self.background_subsection
                    ),  # get_rfp_domain(req),
                    "requirement_domain": domain.classify(
                        req
                    ),  # get_requirement_domain(req),
                    "confidence_factor": get_confidence_factor(req),
                    "relationships": [],
                }
                self.extracted_requirements_list.append(requirement_data)

                # Identifier les exigences liées à la demande actuelle
                requirement_data["relationships"] = identify_relationships(
                    requirement_data, self.all_requirements_data
                )

                # Ajouter les données d'exigence à la liste de toutes les exigences
                self.all_requirements_data.append(requirement_data)

        # Convertissez la liste en JSON et enregistrez-la dans un fichier
        with open("extracted_requirements_cls.json", "w") as json_file:
            json.dump(self.extracted_requirements_list, json_file, indent=4)

    def process_file(self):
        extension = self.file_path.split(".")[-1]
        if extension == "pdf":
            self.read_pdf()
        elif extension == "docx":
            self.read_docx()
        elif extension == "txt":
            self.read_txt()
        else:
            print("Extension de fichier non prise en charge.")
            # return
        self.clean_text()
        self.extract_background_subsection()

        self.paragraphs = self.segment_paragraphs()

        for paragraph in self.paragraphs:
            sentences = self.segment_sentences(paragraph)
            self.segmented_text.extend(sentences)

        self.extract_requirements()


# Utilisation de la classe
file_path = "/Users/mamoutou.doumbia/Desktop/ComplianceCheck/research/data/texteN2.txt"  # Remplacez par le chemin vers votre fichier
extractor = RequirementsExtractor(file_path)
extractor.process_file()
